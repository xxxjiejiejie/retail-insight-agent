from __future__ import annotations

import hashlib
import json
import re
from dataclasses import replace
from pathlib import Path
from typing import Any

from app.rag.models import DocumentChunk, DocumentSection, PolicyDocument

REQUIRED_METADATA = {"document_id", "title", "version", "effective_date"}
HEADING_PATTERN = re.compile(r"^(#{2,4})\s+(.+?)\s*$")
SENTENCE_ENDINGS = "。！？；\n"
SUPPORTED_SUFFIXES = {".md", ".pdf", ".docx"}


def _parse_frontmatter(text: str, source: Path) -> tuple[dict[str, str], str]:
    if not text.startswith("---\n"):
        raise ValueError(f"制度文档缺少 YAML frontmatter：{source.name}")
    try:
        raw_metadata, body = text[4:].split("\n---\n", 1)
    except ValueError as exc:
        raise ValueError(f"制度文档 frontmatter 未闭合：{source.name}") from exc

    metadata: dict[str, str] = {}
    for line in raw_metadata.splitlines():
        key, separator, value = line.partition(":")
        if not separator or not key.strip() or not value.strip():
            raise ValueError(f"制度文档元数据格式错误：{source.name}")
        metadata[key.strip()] = value.strip().strip('"')

    missing = sorted(REQUIRED_METADATA - metadata.keys())
    if missing:
        raise ValueError(f"制度文档缺少元数据 {', '.join(missing)}：{source.name}")
    return metadata, body.strip()


def _load_markdown_document(path: Path) -> PolicyDocument:
    metadata, body = _parse_frontmatter(path.read_text(encoding="utf-8"), path)
    return PolicyDocument(
        document_id=metadata["document_id"],
        title=metadata["title"],
        version=metadata["version"],
        effective_date=metadata["effective_date"],
        source=path.name,
        content=body,
    )


def _load_sidecar_metadata(path: Path) -> dict[str, str]:
    sidecar = path.with_suffix(f"{path.suffix}.metadata.json")
    if not sidecar.exists():
        raise ValueError(f"{path.name} 缺少元数据侧车文件：{sidecar.name}")
    try:
        raw: Any = json.loads(sidecar.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise ValueError(f"元数据侧车文件不是有效 JSON：{sidecar.name}") from exc
    if not isinstance(raw, dict):
        raise ValueError(f"元数据侧车文件必须是 JSON 对象：{sidecar.name}")
    metadata = {str(key): str(value).strip() for key, value in raw.items()}
    missing = sorted(key for key in REQUIRED_METADATA if not metadata.get(key))
    if missing:
        raise ValueError(f"元数据侧车文件缺少字段 {', '.join(missing)}：{sidecar.name}")
    return metadata


def _load_pdf_document(path: Path) -> PolicyDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("读取 PDF 需要安装 pypdf，请安装项目的 rag 可选依赖") from exc

    metadata = _load_sidecar_metadata(path)
    reader = PdfReader(str(path))
    sections: list[DocumentSection] = []
    for page_number, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            sections.append(
                DocumentSection(title=f"第 {page_number} 页", content=text, page=page_number)
            )
    if not sections:
        raise ValueError(f"{path.name} 未提取到文本，可能是扫描版 PDF，需要先进行 OCR")
    return PolicyDocument(
        document_id=metadata["document_id"],
        title=metadata["title"],
        version=metadata["version"],
        effective_date=metadata["effective_date"],
        source=path.name,
        content="\n\n".join(section.content for section in sections),
        sections=tuple(sections),
    )


def _load_docx_document(path: Path) -> PolicyDocument:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError("读取 DOCX 需要安装 python-docx，请安装项目的 rag 可选依赖") from exc

    metadata = _load_sidecar_metadata(path)
    document = Document(str(path))
    sections: list[DocumentSection] = []
    current_title = "总则"
    current_lines: list[str] = []

    def flush_section() -> None:
        content = "\n".join(line for line in current_lines if line).strip()
        if content:
            sections.append(DocumentSection(title=current_title, content=content))

    for block in document.iter_inner_content():
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            if block.style and block.style.name.lower().startswith("heading"):
                flush_section()
                current_title = text
                current_lines = []
            else:
                current_lines.append(text)
        elif isinstance(block, Table):
            for row in block.rows:
                values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(values):
                    current_lines.append(" | ".join(values))
    flush_section()
    if not sections:
        raise ValueError(f"{path.name} 未提取到可检索文本")
    return PolicyDocument(
        document_id=metadata["document_id"],
        title=metadata["title"],
        version=metadata["version"],
        effective_date=metadata["effective_date"],
        source=path.name,
        content="\n\n".join(section.content for section in sections),
        sections=tuple(sections),
    )


def load_policy_document(path: Path) -> PolicyDocument:
    suffix = path.suffix.lower()
    if suffix == ".md":
        return _load_markdown_document(path)
    if suffix == ".pdf":
        return _load_pdf_document(path)
    if suffix == ".docx":
        return _load_docx_document(path)
    raise ValueError(f"不支持的制度文档格式：{path.suffix or path.name}")


def discover_policy_paths(directory: Path) -> list[Path]:
    return sorted(
        path
        for path in directory.rglob("*")
        if path.is_file()
        and path.suffix.lower() in SUPPORTED_SUFFIXES
        and path.name.lower() != "readme.md"
        and not path.name.startswith("~$")
    )


def load_policy_documents(directory: Path) -> list[PolicyDocument]:
    paths = discover_policy_paths(directory)
    documents = [
        replace(load_policy_document(path), source=path.relative_to(directory).as_posix())
        for path in paths
    ]
    identifiers = [document.document_id for document in documents]
    if len(identifiers) != len(set(identifiers)):
        raise ValueError("制度文档 document_id 不能重复")
    return documents


def policy_document_sections(document: PolicyDocument) -> list[DocumentSection]:
    if document.sections:
        return list(document.sections)
    return [
        DocumentSection(title=title, content=content)
        for title, content in _section_blocks(document.content)
    ]


def _section_blocks(content: str) -> list[tuple[str, str]]:
    current_section = "总则"
    current_lines: list[str] = []
    blocks: list[tuple[str, str]] = []
    for line in content.splitlines():
        heading = HEADING_PATTERN.match(line)
        if heading:
            section_text = "\n".join(current_lines).strip()
            if section_text:
                blocks.append((current_section, section_text))
            current_section = heading.group(2).strip()
            current_lines = []
        elif line.startswith("# "):
            continue
        else:
            current_lines.append(line)
    section_text = "\n".join(current_lines).strip()
    if section_text:
        blocks.append((current_section, section_text))
    return blocks


def _split_text(text: str, max_chars: int, overlap_chars: int) -> list[str]:
    compact = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(compact) <= max_chars:
        return [compact]

    parts: list[str] = []
    start = 0
    while start < len(compact):
        hard_end = min(start + max_chars, len(compact))
        end = hard_end
        if hard_end < len(compact):
            minimum_end = start + max_chars // 2
            candidates = [compact.rfind(mark, minimum_end, hard_end) for mark in SENTENCE_ENDINGS]
            sentence_end = max(candidates)
            if sentence_end >= minimum_end:
                end = sentence_end + 1
        part = compact[start:end].strip()
        if part:
            parts.append(part)
        if end >= len(compact):
            break
        start = max(end - overlap_chars, start + 1)
    return parts


def chunk_policy_document(
    document: PolicyDocument,
    *,
    max_chars: int = 700,
    overlap_chars: int = 80,
) -> list[DocumentChunk]:
    if max_chars < 200:
        raise ValueError("max_chars 不能小于 200")
    if overlap_chars < 0 or overlap_chars >= max_chars // 2:
        raise ValueError("overlap_chars 必须大于等于 0 且小于 max_chars 的一半")

    chunks: list[DocumentChunk] = []
    source_sections = [
        (section.title, section.content, section.page)
        for section in policy_document_sections(document)
    ]
    for section_index, (section, section_text, page) in enumerate(source_sections, 1):
        for part_index, part in enumerate(
            _split_text(section_text, max_chars, overlap_chars),
            1,
        ):
            paragraph_id = f"{document.document_id}-S{section_index:02d}-P{part_index:02d}"
            digest_input = f"{document.document_id}|{section}|{paragraph_id}|{part}"
            chunk_id = hashlib.sha256(digest_input.encode("utf-8")).hexdigest()[:20]
            chunks.append(
                DocumentChunk(
                    chunk_id=chunk_id,
                    document_id=document.document_id,
                    title=document.title,
                    version=document.version,
                    effective_date=document.effective_date,
                    source=document.source,
                    section=section,
                    paragraph_id=paragraph_id,
                    content=part,
                    page=page,
                )
            )
    return chunks


def load_and_chunk_policies(
    directory: Path,
    *,
    max_chars: int = 700,
    overlap_chars: int = 80,
) -> list[DocumentChunk]:
    return [
        chunk
        for document in load_policy_documents(directory)
        for chunk in chunk_policy_document(
            document,
            max_chars=max_chars,
            overlap_chars=overlap_chars,
        )
    ]
